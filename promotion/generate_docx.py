from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_everytime_post():
    doc = Document()

    # Title
    title = doc.add_heading('[에브리타임 홍보 게시글 초안]', 0)
    
    # Header Note
    p = doc.add_paragraph()
    run = p.add_run('※ 아래 내용을 복사하여 에브리타임 커뮤니티(컴공/공대 게시판)에 게시하세요.')
    run.font.size = Pt(10)
    run.italic = True

    doc.add_paragraph('---')

    # Main Title
    p_main_title = doc.add_paragraph()
    run_main_title = p_main_title.add_run('제목: [도와주세요] 번아웃 연구하는 컴공생입니다... 1분만 빌려주세요! 😭')
    run_main_title.bold = True
    run_main_title.font.size = Pt(14)

    # Body
    doc.add_paragraph('\n안녕하세요, 개발자 번아웃과 코딩 패턴의 상관관계를 주제로 연구 중인 학생입니다.')
    doc.add_paragraph('코딩하다 보면 "내가 지금 잘하고 있나?", "갑자기 현타 온다" 싶을 때가 있죠? 이런 주관적인 느낌이 실제로 타이핑 속도나 에러 해결 패턴에 어떻게 나타나는지 데이터로 증명해 보려고 합니다.')
    doc.add_paragraph('한 명의 간절한 공대생을 구제해 주실 분들을 찾습니다!')

    # Section: How to participate
    p_steps = doc.add_paragraph()
    run_steps = p_steps.add_run('✅ 참여 방법 (딱 1분 걸려요!)')
    run_steps.bold = True
    
    doc.add_paragraph('1. VS Code 확장 프로그램 탭에서 "Burnout Detector" 검색 후 설치', style='List Number')
    doc.add_paragraph('2. 설치 직후 뜨는 알림창에서 [동의함(Agree)] 클릭', style='List Number')
    doc.add_paragraph('3. 끝! 평소처럼 과제나 프로젝트 코딩하시면 됩니다.', style='List Number')

    # Section: Security
    p_sec = doc.add_paragraph()
    run_sec = p_sec.add_run('\n🔒 보안 및 개인정보 (가장 중요!)')
    run_sec.bold = True

    doc.add_paragraph('• 코드 내용(Text)은 절대 수집하지 않습니다. (보안 걱정 NO!)', style='List Bullet')
    doc.add_paragraph('• 오직 행동 지표(타이핑 빈도, 유휴 시간, 에러 해결 소요 시간 등)만 기록됩니다.', style='List Bullet')
    doc.add_paragraph('• 모든 데이터는 100% 익명화되어 처리되니 안심하셔도 됩니다.', style='List Bullet')

    # Section: Promise
    p_promise = doc.add_paragraph()
    run_promise = p_promise.add_run('\n🎁 연구 결과 공유')
    run_promise.bold = True
    doc.add_paragraph('여러분의 소중한 데이터로 의미 있는 결과가 나오면, 나중에 에타에 분석 리포트를 요약해서 공유하겠습니다! (우리 과 학생들의 평균 코딩 집중 시간은 얼마나 될까? 등)')

    doc.add_paragraph('\n여러분의 데이터 한 조각이 저의 연구와 졸업에 정말 큰 힘이 됩니다. 미리 감사드립니다! 🙇‍♂️')
    
    # Link
    doc.add_paragraph('\n마켓플레이스 링크: https://marketplace.visualstudio.com/items?itemName=hojun-lee.burnout-detector')

    # Save
    doc.save('C:/Users/이호준/burnout-detector/에브리타임_홍보_게시글.docx')
    print('Successfully created the docx file.')

if __name__ == '__main__':
    create_everytime_post()
